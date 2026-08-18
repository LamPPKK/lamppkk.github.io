#!/usr/bin/env python3
"""Rebuild the copied reference CV with canonical portfolio data and semantic styles."""

from pathlib import Path
from datetime import datetime, timezone
import json
import shutil
import subprocess
import sys
import zipfile
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ORANGE = "D4551C"
NAVY = "071B33"
MUTED = "526273"
LIGHT = "EEE9DD"
CANONICAL_NAME = "Nguyễn Đức Tùng Lâm (Liam)"
SNAPSHOT_TIME = datetime(2026, 8, 19, tzinfo=timezone.utc)


def set_cell_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def set_bottom_border(paragraph, color=ORANGE, size="10"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_hyperlink(paragraph, text, url, color="1D5FD1"):
    relation = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([run_color, underline])
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(7.5)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def ensure_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_nums = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abstract, default=-1) + 1
    num_id = max(existing_nums, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    paragraph_properties = OxmlElement("w:pPr")
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "245")
    indentation.set(qn("w:hanging"), "173")
    paragraph_properties.append(indentation)
    level.extend([start, number_format, level_text, justification, paragraph_properties])
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def apply_bullet(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    p_pr.insert(0, num_pr)


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(text, style="Heading 1")
    set_bottom_border(paragraph)
    return paragraph


def add_role(doc, item, bullet_num_id, compact=False):
    title = f"{item['role']['en']} — {item['employer']}"
    doc.add_paragraph(title, style="Heading 2")
    end = item.get("end_date") or "Present"
    meta = doc.add_paragraph(style="CV Meta")
    meta.add_run(f"{item['start_date']} — {end}  ·  {item['employment_type']['en']}  ·  {item['location']['en']}")
    bullet = doc.add_paragraph(item["summary"]["en"], style="List Bullet")
    apply_bullet(bullet, bullet_num_id)
    if compact:
        bullet.paragraph_format.space_after = Pt(2)


def add_project(doc, item, bullet_num_id):
    doc.add_paragraph(item["title"]["en"], style="Heading 2")
    meta = doc.add_paragraph(style="CV Meta")
    period = item["period"]["en"] if isinstance(item["period"], dict) else item["period"]
    meta.add_run(f"{item['role']['en']}  ·  {period}")
    challenge = doc.add_paragraph(item["challenge"]["en"], style="List Bullet")
    contribution = doc.add_paragraph(item["contribution"]["en"][0], style="List Bullet")
    apply_bullet(challenge, bullet_num_id)
    apply_bullet(contribution, bullet_num_id)
    tech = doc.add_paragraph(style="CV Tech")
    tech.add_run("Technology: ").bold = True
    tech.add_run(" · ".join(item["technologies"]))
    if item["privacy"] == "private-anonymized":
        note = doc.add_paragraph("Private engagement — client name and repository withheld.", style="CV Note")
        note.paragraph_format.space_after = Pt(4)
    elif item.get("links"):
        note = doc.add_paragraph(style="CV Note")
        note.add_run("Public code: ")
        for index, link in enumerate(item["links"]):
            if index:
                note.add_run(" · ")
            add_hyperlink(note, link["url"].rstrip("/").split("/")[-1], link["url"])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Liberation Sans"
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04
    normal.paragraph_format.widow_control = True

    for name, size, color in (("Title", 25, NAVY), ("Subtitle", 10.5, ORANGE), ("Heading 1", 12.5, NAVY), ("Heading 2", 9.5, NAVY)):
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Liberation Sans"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True
    styles["Title"].paragraph_format.space_after = Pt(1)
    styles["Subtitle"].paragraph_format.space_after = Pt(6)
    styles["Heading 1"].paragraph_format.space_before = Pt(8)
    styles["Heading 1"].paragraph_format.space_after = Pt(5)
    styles["Heading 2"].paragraph_format.space_before = Pt(4)
    styles["Heading 2"].paragraph_format.space_after = Pt(1)
    bullet_style = styles["List Bullet"] if "List Bullet" in styles else styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = "Liberation Sans"
    bullet_style.font.size = Pt(8.4)
    bullet_style.paragraph_format.left_indent = Inches(0.17)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.12)
    bullet_style.paragraph_format.space_after = Pt(2)
    bullet_style.paragraph_format.widow_control = True

    custom = {
        "CV Contact": (8.3, NAVY, False),
        "CV Summary": (9.3, NAVY, False),
        "CV Meta": (7.3, MUTED, False),
        "CV Tech": (7.5, MUTED, False),
        "CV Note": (7.2, "14715B", True),
        "CV Credential": (8.1, NAVY, False),
    }
    for name, (size, color, italic) in custom.items():
        if name not in styles:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[name]
        style.base_style = styles["Normal"]
        style.font.name = "Liberation Sans"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.italic = italic
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.widow_control = True
    styles["CV Summary"].paragraph_format.space_after = Pt(5)
    styles["CV Summary"].paragraph_format.line_spacing = 1.08


def sanitize_package_metadata(output_path):
    """Remove template-owner metadata that can survive a normal python-docx save."""
    temporary_path = output_path.with_suffix(".sanitized.docx")

    with zipfile.ZipFile(output_path, "r") as source, zipfile.ZipFile(temporary_path, "w") as target:
        for info in source.infolist():
            if info.filename in ("docProps/custom.xml", "docMetadata/LabelInfo.xml"):
                continue
            data = source.read(info.filename)
            if info.filename == "docProps/app.xml":
                text = data.decode("utf-8")
                values = {
                    "Template": "Normal.dotm",
                    "TotalTime": "0",
                    "Pages": "3",
                    "Words": "0",
                    "Characters": "0",
                    "Application": "Microsoft Office Word",
                    "Lines": "0",
                    "Paragraphs": "0",
                    "Company": "",
                    "Manager": "",
                    "CharactersWithSpaces": "0",
                }
                for field, value in values.items():
                    text = re.sub(fr"(<{field}>).*?(</{field}>)", fr"\g<1>{value}\g<2>", text, flags=re.DOTALL)
                text = text.replace("<vt:lpstr>CV Template</vt:lpstr>", "<vt:lpstr>Curriculum Vitae</vt:lpstr>")
                data = text.encode("utf-8")
            elif info.filename == "_rels/.rels":
                data = re.sub(rb"<Relationship\b[^>]*(?:custom-properties|classificationlabels)[^>]*/>", b"", data)
            elif info.filename == "[Content_Types].xml":
                data = re.sub(rb"<Override\b[^>]*PartName=[\"']/(?:docProps/custom|docMetadata/LabelInfo)\.xml[\"'][^>]*/>", b"", data)
            target.writestr(info, data)
    temporary_path.replace(output_path)


def main(source_path, output_path):
    data = ROOT / "_data"
    def load_yaml(path):
        result = subprocess.run(
            ["ruby", "-ryaml", "-rjson", "-e", "print JSON.generate(YAML.load_file(ARGV[0]))", str(path)],
            check=True,
            capture_output=True,
            text=True,
        )
        return json.loads(result.stdout)

    profile = load_yaml(data / "profile.yml")
    experience = load_yaml(data / "experience.yml")
    projects = load_yaml(data / "projects.yml")
    credentials = load_yaml(data / "credentials.yml")
    skills = load_yaml(data / "skills.yml")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if source_path.resolve() != output_path.resolve():
        shutil.copyfile(source_path, output_path)
    doc = Document(output_path)
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.46)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)
    doc.settings.odd_and_even_pages_header_footer = False
    settings = doc.settings.element
    for name in ("removePersonalInformation", "removeDateAndTime"):
        if settings.find(qn(f"w:{name}")) is None:
            settings.append(OxmlElement(f"w:{name}"))
    section.different_first_page_header_footer = False
    configure_styles(doc)
    bullet_num_id = ensure_bullet_numbering(doc)

    core = doc.core_properties
    core.title = "Nguyễn Đức Tùng Lâm (Liam) — Technical Project Manager CV"
    core.subject = "Professional experience, selected projects, skills, education, and credentials"
    core.author = CANONICAL_NAME
    core.last_modified_by = CANONICAL_NAME
    core.created = SNAPSHOT_TIME
    core.modified = SNAPSHOT_TIME
    last_printed = core._element.find("{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastPrinted")
    if last_printed is not None:
        core._element.remove(last_printed)
    core.revision = 1
    core.category = "Professional portfolio"
    core.content_status = "Published"
    core.identifier = "lamppkk-cv-2026-08-19"
    core.language = "en"
    core.version = "1.0"
    core.keywords = "Technical Project Manager, iOS, mobile, security assessment, software operations"
    core.comments = f"Canonical profile snapshot: {profile['snapshot_date']}"

    for header_part in (section.header, section.first_page_header, section.even_page_header):
        header = header_part.paragraphs[0]
        header.text = "NGUYỄN ĐỨC TÙNG LÂM  ·  TECHNICAL PROJECT MANAGEMENT"
        header.style = "CV Meta"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(section.footer.paragraphs[0])

    accent = doc.add_paragraph(" ")
    accent.paragraph_format.space_after = Pt(5)
    accent.paragraph_format.line_spacing = .25
    set_cell_shading(accent, ORANGE)
    doc.add_paragraph(profile["name"], style="Title")
    doc.add_paragraph(profile["headline"]["en"], style="Subtitle")
    contact = doc.add_paragraph(style="CV Contact")
    contact.add_run(f"{profile['location']['en']}  ·  {profile['phone']}  ·  ")
    add_hyperlink(contact, profile["email"], f"mailto:{profile['email']}")
    contact.add_run("  ·  ")
    add_hyperlink(contact, "LinkedIn", profile["links"]["linkedin"])
    contact.add_run("  ·  ")
    add_hyperlink(contact, "GitHub", profile["links"]["github"])
    contact.add_run("  ·  ")
    add_hyperlink(contact, "Portfolio", "https://lamppkk.github.io/")

    add_section_heading(doc, "Professional profile")
    doc.add_paragraph(profile["about"]["en"][0], style="CV Summary")
    doc.add_paragraph(profile["about"]["en"][1], style="CV Summary")

    add_section_heading(doc, "Core expertise")
    for group in skills[:4]:
        paragraph = doc.add_paragraph(style="CV Tech")
        run = paragraph.add_run(f"{group['name']['en']}: ")
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(ORANGE)
        paragraph.add_run(" · ".join(group["items"]["en"]))

    add_section_heading(doc, "Leadership & operations experience")
    first_page_ids = [
        "fpt-software-project-manager-2026-06",
        "fpt-software-associate-pm-2025-06",
        "fpt-software-ios-developer-2021-08",
        "viettel-cyber-security-pm-2026-04",
        "nodesign-operations-pm-2024-06",
        "viettel-software-junior-pm-2026-02",
        "fpt-polytechnic-it-lecturer-2024-02",
        "fpt-software-innovation-mobile-engineer-2022-03",
        "fpt-software-academy-fresher-ios-2021-03",
    ]
    by_id = {item["id"]: item for item in experience}
    for item_id in first_page_ids:
        add_role(doc, by_id[item_id], bullet_num_id, compact=True)

    doc.add_page_break()
    add_section_heading(doc, "Engineering, education & early experience")
    second_page_ids = [item["id"] for item in experience if item["id"] not in first_page_ids]
    for item_id in second_page_ids:
        add_role(doc, by_id[item_id], bullet_num_id, compact=True)

    add_section_heading(doc, "Selected delivery & security projects")
    project_by_id = {item["id"]: item for item in projects}
    for item_id in ["enterprise-mobile-security", "japanese-banking-security", "creative-operations-platform"]:
        add_project(doc, project_by_id[item_id], bullet_num_id)
    for item_id in ["multi-market-mobile-banking", "luxury-ar-commerce"]:
        add_project(doc, project_by_id[item_id], bullet_num_id)

    doc.add_page_break()
    add_section_heading(doc, "Selected mobile & open-source projects")
    for item_id in ["airport-operations-mobile", "academic-mobile-portal", "xanhtab", "fireball-browser-systems", "smart-movie", "waydroid-wsl"]:
        add_project(doc, project_by_id[item_id], bullet_num_id)

    add_section_heading(doc, "Education")
    education = profile["education"][0]
    paragraph = doc.add_paragraph(style="CV Credential")
    paragraph.add_run(education["school"] + " — ").bold = True
    paragraph.add_run(f"{education['degree']['en']}  ·  {education['start_date']} — {education['end_date']}")

    add_section_heading(doc, "Selected credentials")
    featured = [item for item in credentials if item.get("featured")][:8]
    for item in featured:
        paragraph = doc.add_paragraph(style="CV Credential")
        paragraph.add_run(item["title"]["en"] + " — ").bold = True
        paragraph.add_run(f"{item['issuer']} · {item['issued']}")
    credentials_link = doc.add_paragraph(style="CV Contact")
    credentials_link.add_run("Complete credential ledger: ")
    add_hyperlink(credentials_link, "lamppkk.github.io/credentials/", "https://lamppkk.github.io/credentials/")

    add_section_heading(doc, "Languages")
    language = doc.add_paragraph(style="CV Credential")
    language.add_run("  ·  ".join(f"{item['name']['en']}: {item['level']['en']}" for item in profile["languages"]))

    doc.save(output_path)
    sanitize_package_metadata(output_path)
    print(f"Built semantic CV: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_cv.py SOURCE_TEMPLATE.docx OUTPUT.docx")
    main(Path(sys.argv[1]), Path(sys.argv[2]))
